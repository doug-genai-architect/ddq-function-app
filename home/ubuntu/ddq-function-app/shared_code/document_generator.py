# shared_code/document_generator.py

import os
import logging
import io
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .blob_storage_service import BlobStorageService # Relative import


def generate_and_upload_docx(question: str, answer: str, sources: list, blob_service: BlobStorageService, 
                      blob_prefix="ddq_responses", template_name=None) -> str | None:
    """Generates a DOCX document with the question, answer, and sources, and uploads it to Blob Storage.

    Args:
        question: The user's DDQ question.
        answer: The AI-generated answer.
        sources: A list of source file names.
        blob_service: An initialized instance of BlobStorageService.
        blob_prefix: The prefix for the blob name in the container.
        template_name: Optional name of a document template to use.

    Returns:
        The URL of the uploaded document, or None if upload fails.
    """
    try:
        logging.info(f"Generating DOCX document {'using template: ' + template_name if template_name else '...'}")
        
        # --- Initialize Document (with template if specified) ---
        templates_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "templates"))
        if template_name and os.path.exists(os.path.join(templates_dir, f"{template_name}.docx")):
            template_path = os.path.join(templates_dir, f"{template_name}.docx")
            document = Document(template_path)
            logging.info(f"Using template from: {template_path}")
        else:
            document = Document()
            # Apply default formatting
            style = document.styles["Normal"]
            font = style.font
            font.name = "Calibri"
            font.size = Pt(11)
            
            # --- Title ---
            title = document.add_heading("Due Diligence Question Response", level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph() # Add some space
            
            # --- Question Section ---
            document.add_heading("Question:", level=2)
            q_paragraph = document.add_paragraph()
            q_paragraph.add_run(question).bold = True
            document.add_paragraph()
            
            # --- Answer Section ---
            document.add_heading("Answer:", level=2)
            document.add_paragraph(answer)
            document.add_paragraph()
            
            # --- Sources Section ---
            document.add_heading("Sources Consulted:", level=2)
            if sources:
                for source in sources:
                    document.add_paragraph(source, style="List Bullet")
            else:
                document.add_paragraph("No specific documents were cited for this response.")
        
        # --- Add Metadata to Document Properties ---
        # This is useful for tracking and searching documents later
        core_properties = document.core_properties
        core_properties.title = f"DDQ Response: {question[:50]}..." if len(question) > 50 else f"DDQ Response: {question}"
        core_properties.subject = "Due Diligence Questionnaire"
        core_properties.created = datetime.datetime.now()
        core_properties.modified = datetime.datetime.now()
        core_properties.category = "DDQ Responses"
        core_properties.keywords = "DDQ, Due Diligence, AI Response"
        
        # --- Save document to memory ---
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0) # Reset stream position to the beginning
        document_content = file_stream.read()
        logging.info("DOCX document generated successfully in memory.")

        # --- Create a meaningful filename ---
        # Generate a safe filename from the question (remove invalid characters)
        safe_filename = "".join([c for c in question[:30] if c.isalnum() or c in " _-"]).strip()
        safe_filename = safe_filename.replace(" ", "_")
        
        # Add a timestamp and random component for uniqueness
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        random_suffix = os.urandom(2).hex()
        
        # Construct the final blob name
        blob_name = f"{blob_prefix}/{safe_filename}_{timestamp}_{random_suffix}.docx"
        
        # --- Upload to Blob Storage ---
        upload_url = blob_service.upload_document(
            document_content=document_content,
            blob_name=blob_name,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        if upload_url:
            logging.info(f"Document successfully uploaded to Blob Storage: {upload_url}")
            return upload_url
        else:
            logging.error("Failed to upload generated DOCX document to Blob Storage.")
            return None

    except Exception as e:
        logging.error(f"Error generating or uploading DOCX document: {e}", exc_info=True)
        return None

# Example usage (for testing purposes - requires BlobStorageService to be configured)
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    try:
        # This requires environment variables for BlobStorageService to be set
        print("Attempting to initialize Blob Storage Service for doc gen test...")
        test_blob_service = BlobStorageService()
        
        print("\nGenerating and uploading example document...")
        test_question = "What is the investment strategy?"
        test_answer = "The investment strategy focuses on acquiring distressed assets and maximizing value through active management."
        test_sources = ["LSREF VII - Amended and Restated PPM November 2022 - Final.docx", "Lone Star Overview - November 2024.pdf"]
        
        doc_url = generate_and_upload_docx(test_question, test_answer, test_sources, test_blob_service)
        
        if doc_url:
            print(f"\nExample document generated and uploaded successfully!")
            print(f"URL: {doc_url}")
        else:
            print("\nFailed to generate or upload example document.")
            
    except ValueError as e:
        print(f"Configuration Error: {e}")
        print("Please ensure Azure Blob Storage environment variables are set.")
    except Exception as e:
        print(f"An unexpected error occurred during example run: {e}")

